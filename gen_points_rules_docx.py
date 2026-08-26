# -*- coding: utf-8 -*-
"""
生成《MOYUYO 会员积分规则说明》Word 文档
规则内容整理自项目源码与需求文档：
- moyuyo-server 积分相关 Controller / Service / Entity / SQL 迁移脚本
- docs/需求文档.md
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x00, 0x5A, 0xB8)      # 主题蓝
HEADER_BG = "D9E8F7"                       # 表头浅蓝
ALT_BG = "F4F8FC"                          # 隔行浅灰蓝
TEXT = RGBColor(0x33, 0x33, 0x33)


def set_cn_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_cn_font(run, name="黑体", size=15, bold=True, color=PRIMARY)
        # 标题下加一条细线
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "005AB8")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        set_cn_font(run, name="黑体", size=12.5, bold=True, color=PRIMARY)
    else:
        set_cn_font(run, name="黑体", size=11, bold=True, color=TEXT)


def add_para(doc, text, indent=0, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, indent=0.5, symbol="• "):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    run = p.add_run(symbol + text)
    set_cn_font(run, size=10.5)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_cn_font(run, name="黑体", size=10, bold=True)
        shade_cell(cell, HEADER_BG)
    # 数据行
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(v))
            set_cn_font(run, size=10)
            if i % 2 == 1:
                shade_cell(cell, ALT_BG)
    if widths:
        for j, w in enumerate(widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    # 段落间距
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


doc = Document()

# 全局页面设置
for section in doc.sections:
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

# 默认字体
style = doc.styles["Normal"]
style.font.name = "宋体"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.font.size = Pt(10.5)

# ============ 封面标题 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("MOYUYO 宠物用品商城")
set_cn_font(run, name="黑体", size=22, bold=True, color=PRIMARY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("会员积分规则说明")
set_cn_font(run, name="黑体", size=28, bold=True, color=TEXT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("（V1.0 · 整理自项目源码与需求文档）")
set_cn_font(run, size=12, color=RGBColor(0x88, 0x88, 0x88))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
run = p.add_run("整理日期：2026-08-25")
set_cn_font(run, size=12, color=RGBColor(0x88, 0x88, 0x88))

# ============ 一、概述 ============
add_heading(doc, "一、积分体系概述", 1)
add_para(doc, "积分是 MOYUYO 商城面向会员用户的虚拟奖励资产，与成长值（会员等级）相互独立。"
              "用户通过签到、购物、邀请、评价、做任务等行为获得积分，可在下单时抵扣现金、兑换礼品或参与抽奖等。")
add_bullet(doc, "积分账户：积分余额记录在用户表（mo_user.points）中，每次变动同时写入积分流水表（mo_points_log）。")
add_bullet(doc, "积分核算：余额 = 累计获得 − 累计消耗，积分不足时系统拒绝扣减并提示“积分不足”。")
add_bullet(doc, "积分与成长值相互独立：成长值决定会员等级（L1~L5），积分决定可兑换/抵扣的额度，两者获取规则互不相同。")
add_bullet(doc, "会员等级对应积分倍率加成（见第五章），等级越高，购物等场景获得的积分越多。")

# ============ 二、积分获取 ============
add_heading(doc, "二、积分获取规则", 1)

add_heading(doc, "2.1 每日签到", 2)
add_bullet(doc, "每日签到 +5 积分（接口 POST /api/v1/points/checkin，流水类型 CHECKIN）。")
add_bullet(doc, "连续签到 7 天，积分奖励 ×2 倍率；签到日历按月展示，可查看历史签到记录。")
add_bullet(doc, "补签规则：每月可免费补签 1 次，之后每次补签消耗 50 积分。")
add_bullet(doc, "当天重复签到会被拒绝（同一用户同一天仅可签到一次）。")

add_heading(doc, "2.2 购物返积分", 2)
add_bullet(doc, "订单完成后按消费行为发放积分，流水类型为 PURCHASE（管理端可配置对应活动）。")
add_bullet(doc, "发放倍率受会员等级影响：Member 1.0x、Silver 1.1x、Gold 1.2x、Platinum 1.5x、Black 2.0x。")
add_bullet(doc, "生日当月消费按 ×2 倍率计算积分（需求规划）。")
add_bullet(doc, "订单取消时，已预扣/使用的积分原路退回（需求规划）。")

add_heading(doc, "2.3 邀请好友", 2)
add_bullet(doc, "每邀请 1 位好友注册，邀请人与被邀请人双方各得 200 积分（前端邀请页文案），流水类型 INVITE。")
add_bullet(doc, "邀请状态流转：PENDING → REGISTERED → ORDERED，好友完成首单后发放奖励（需求文档：双方各得 $10 积分奖励）。")
add_bullet(doc, "邀请榜单：月度 TOP 10 邀请者可获得额外积分奖励（需求规划）。")

add_heading(doc, "2.4 商品评价", 2)
add_bullet(doc, "文字评价 +10 积分、含图评价 +20 积分、含视频评价 +5 积分。")
add_bullet(doc, "评价积分每日上限 100 分。")
add_bullet(doc, "订单完成后 7 天未评价推送提醒，15 天未评价自动默认 5 星好评（可修改）。")

add_heading(doc, "2.5 新人任务", 2)
add_bullet(doc, "新人注册引导：完善宠物档案 +50 积分、设置护理提醒 +30 积分、订阅 Pet Hub 每日签到 +20 积分。")
add_bullet(doc, "新人专区注册即得 $5 积分（需求规划，可直接抵现），新人权益有效期为注册后 30 天。")

add_heading(doc, "2.6 任务中心", 2)
add_para(doc, "任务分为每日 / 每周 / 成就三类，任务定义存于 mo_mission 表（type 区分 DAILY / WEEKLY / ACHIEVEMENT，points 为奖励积分），"
              "由后台配置，用户任务进度记录在 mo_user_mission 表。奖励需满足完成条件后手动领取（接口 POST /api/v1/missions/{id}/claim）。")
add_table(doc,
    ["任务类型", "任务内容", "积分奖励"],
    [
        ["每日任务", "每日签到", "+5"],
        ["每日任务", "浏览 5 个商品", "+3"],
        ["每日任务", "分享 1 个商品", "+5"],
        ["每日任务", "完成 1 次 Pet Hub 互动", "+2"],
        ["每周任务", "累计签到 5 天", "+20"],
        ["每周任务", "完成 1 单购物", "+30"],
        ["每周任务", "发布 1 条社区笔记", "+15"],
        ["每周任务", "邀请 1 位好友注册", "+50"],
        ["成就任务", "首单完成", "+100（附新人徽章）"],
        ["成就任务", "累计消费满 $500", "+200（附消费达人徽章）"],
        ["成就任务", "连续签到 30 天", "+100（附签到达人徽章）"],
        ["成就任务", "发布 10 条笔记", "+80（附创作者徽章）"],
        ["成就任务", "邀请 10 位好友", "+500（附邀请达人徽章）"],
    ],
    widths=[3.0, 6.2, 5.2])
add_para(doc, "任务刷新：每日任务每天 0 点刷新，周任务每周一 0 点刷新。任务奖励包含积分、成长值、优惠券、徽章及限定 Pet Hub 场景。", size=10, color=RGBColor(0x66, 0x66, 0x66))

add_heading(doc, "2.7 活动奖励与人工调整", 2)
add_bullet(doc, "运营活动：后台可创建积分活动（类型 EVENT / SIGN_IN / PURCHASE / INVITE），向用户发放积分，流水类型与活动类型一致。")
add_bullet(doc, "人工调整：后台可对指定用户手动增减积分（类型 ADJUST，需填写调整原因），正数为奖励、负数为扣减。")
add_bullet(doc, "抽奖奖励：抽奖奖品中包含积分奖品（需求规划：奖品为优惠券 / 积分 / 实物礼品 / 大奖礼盒）。")

# ============ 三、积分消耗 ============
add_heading(doc, "三、积分消耗规则", 1)

add_heading(doc, "3.1 订单抵现", 2)
add_bullet(doc, "抵扣比例：100 积分 = $1，单笔订单最高抵扣订单金额的 30%。")
add_bullet(doc, "订单记录抵扣字段：points_discount（抵扣金额）、points_used（使用的积分数量）。")
add_bullet(doc, "抵扣校验：下单时系统校验积分余额充足，余额不足则拒绝抵扣。")

add_heading(doc, "3.2 积分商城兑换", 2)
add_bullet(doc, "用户可使用积分兑换商城礼品（如品牌帆布袋 500 积分、$5 优惠券 1000 积分、免邮券 800 积分等，商品列表由前端/后台配置）。")
add_bullet(doc, "兑换实物礼品需另付运费，兑换码 / 券类即时到账；积分不足时无法兑换。")

add_heading(doc, "3.3 抽奖", 2)
add_bullet(doc, "抽奖活动配置于 mo_lottery 表：points_cost 为单次消耗积分，daily_free 为每日免费次数。")
add_bullet(doc, "每日免费次数用完后，再次抽奖需消耗 points_cost 积分，中奖概率由奖品概率表控制。")

add_heading(doc, "3.4 其他消耗", 2)
add_bullet(doc, "漏签补签：免费次数用完后，每次补签消耗 50 积分。")
add_bullet(doc, "兑换 Pet Hub 场景：可用积分兑换互动场景（需求规划，流水类型 EXCHANGE）。")

# ============ 四、有效期 ============
add_heading(doc, "四、积分有效期与退回规则", 1)
add_bullet(doc, "有效期：每笔积分自获得之日起 12 个月有效，到期未使用自动清零（流水类型 EXPIRE，需求规划）。")
add_bullet(doc, "到期提醒：积分到期前 30 天向用户推送提醒（需求规划）。")
add_bullet(doc, "订单退款：退款时已使用的积分不退回；订单未使用但已获得的积分退回，退回后有效期保持不变（需求规划）。")
add_bullet(doc, "订单取消：待支付订单超时取消后，预扣积分原路退回（需求规划）。")

# ============ 五、会员等级 ============
add_heading(doc, "五、会员等级与积分倍率", 1)
add_para(doc, "会员共 5 个等级，按成长值累计升级。成长值与积分相互独立，但等级影响积分获取倍率。")
add_table(doc,
    ["等级", "成长值门槛", "名称", "定位", "积分倍率"],
    [
        ["L1", "0", "Member", "注册即获得", "1.0x"],
        ["L2", "500", "Silver", "完成首单 + 几次签到", "1.1x"],
        ["L3", "2,000", "Gold", "活跃用户", "1.2x"],
        ["L4", "8,000", "Platinum", "高频消费用户", "1.5x"],
        ["L5", "25,000", "Black", "顶级 VIP", "2.0x"],
    ],
    widths=[1.8, 2.6, 2.6, 4.6, 2.4])
add_para(doc, "成长值获取参考：注册 +50、首单 +100、实付消费 1 美元 = 1 点（退款扣减）、完成含图评价 +20/条（上限 100/日）、"
              "晒宠发布 +30/条（上限 90/日）、邀请好友首单 +200/人、每日签到 +5、Pet Hub 互动 +3（上限 9/日）、"
              "完善宠物档案 +30/只、生日当月消费 ×2 倍率。", size=10, color=RGBColor(0x66, 0x66, 0x66))
add_para(doc, "升降级规则：成长值达到门槛即时升级、当月生效；每自然年累计成长值不足该等级门槛 60% 则次年 1 月 1 日降一级，"
              "降级前 30 天预警；Black 会员连续 2 年保级成功后第 3 年自动保级。", size=10, color=RGBColor(0x66, 0x66, 0x66))

# ============ 六、技术实现 ============
add_heading(doc, "六、技术实现（数据表与接口）", 1)

add_heading(doc, "6.1 数据表", 2)
add_table(doc,
    ["表名", "说明", "关键字段"],
    [
        ["mo_user", "用户表（积分账户）", "points：当前积分余额"],
        ["mo_points_log", "积分流水表", "user_id、change_value（正负值）、type、biz_no、remark、created_at、deleted"],
        ["mo_mission", "任务定义表", "type（DAILY/WEEKLY/ACHIEVEMENT）、points（奖励积分）、target、active"],
        ["mo_user_mission", "用户任务进度表", "progress、completed、claimed"],
        ["mo_invite", "邀请记录表", "invite_code、invited_user_id、status（PENDING/REGISTERED/ORDERED）、points_awarded"],
        ["mo_lottery", "抽奖活动表", "points_cost（单次消耗积分）、daily_free（每日免费次数）、probability"],
        ["mo_lottery_record", "抽奖记录表", "used_free_spin、points_spent、won、prize_name"],
        ["mo_order", "订单表", "points_discount（积分抵扣金额）、points_used（使用积分数）"],
    ],
    widths=[3.4, 5.4, 5.6])

add_heading(doc, "6.2 积分流水类型", 2)
add_table(doc,
    ["类型", "含义", "方向"],
    [
        ["CHECKIN", "每日签到", "获取"],
        ["PURCHASE / ORDER", "购物返积分", "获取"],
        ["INVITE", "邀请奖励", "获取"],
        ["EVENT", "活动奖励", "获取"],
        ["ADJUST", "后台人工调整", "获取 / 扣减"],
        ["SIGN_IN", "签到活动（管理端活动类型）", "获取"],
        ["SPEND", "积分消费", "扣减"],
        ["EXCHANGE", "积分兑换", "扣减"],
        ["EXPIRE", "积分过期清零", "扣减"],
        ["REFUND", "退款退回", "获取"],
    ],
    widths=[3.0, 7.4, 4.0])

add_heading(doc, "6.3 相关接口", 2)
add_table(doc,
    ["接口", "说明"],
    [
        ["GET /api/v1/points/log", "查询本人积分流水（分页）"],
        ["GET /api/v1/points/balance", "查询本人积分余额"],
        ["POST /api/v1/points/checkin", "每日签到（+5 积分）"],
        ["GET /api/v1/missions", "任务列表（含奖励积分）"],
        ["GET /api/v1/missions/my", "我的任务状态"],
        ["POST /api/v1/missions/{id}/claim", "领取任务奖励"],
        ["GET /api/v1/missions/stats", "任务统计（今日已获积分 / 连续签到）"],
        ["GET /api/admin/points/activities", "管理端：积分活动列表"],
        ["POST /api/admin/points/activities/create", "管理端：创建积分活动"],
        ["GET /api/admin/points/logs", "管理端：积分流水查询"],
        ["GET /api/admin/points/stats", "管理端：积分统计（发放/消耗/参与用户）"],
        ["POST /api/admin/points/users/{userId}/adjust", "管理端：手动调整用户积分"],
    ],
    widths=[8.0, 6.4])

# ============ 七、说明 ============
add_heading(doc, "七、规则来源与补充说明", 1)
add_bullet(doc, "已落地实现（源码可直接验证）：每日签到 +5 积分、积分余额 / 流水查询、任务中心领取奖励、管理端积分活动与人工调整、邀请记录与抽奖数据结构。")
add_bullet(doc, "需求规划中（docs/需求文档.md，部分尚未在业务代码中完全打通）：订单抵现（100 积分 = $1、最高抵 30%）、积分 12 个月有效期、退款积分退回、评价奖励、新人奖励、邀请首单 $10 积分奖励、生日双倍积分。")
add_bullet(doc, "前端邀请页文案为“双方各得 200 积分”，需求文档为“好友首单后双方各得 $10 积分”，两者口径暂不一致，上线前需运营确认统一。")
add_bullet(doc, "购物返积分倍率（1.0x~2.0x）与订单抵现参数（100 积分 = $1、抵 30% 上限）当前以配置/常量形式存在于代码与前端，后续建议收敛到后台可配置。")

# 页脚页码
footer_p = doc.sections[0].footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run("MOYUYO 会员积分规则说明 · 第 1 页")
set_cn_font(run, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))

out = r"D:\MOYUYOWPC\MOYUYO会员积分规则说明.docx"
doc.save(out)
print("saved:", out)
